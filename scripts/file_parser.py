"""
CommScribe File Parser
Extracts text from various document formats for voice learning.

Supported formats:
- .txt, .md (plain text)
- .docx (Word 2007+)
- .doc (legacy Word, best-effort)
- .pdf (text-based, not scanned)
- .rtf (rich text)
"""

import os
import re
from pathlib import Path
from typing import List, Tuple, Optional
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    filename: str
    format: str
    text: str
    word_count: int
    success: bool
    error: Optional[str] = None


def parse_txt(filepath: str) -> str:
    """Parse plain text or markdown file."""
    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def parse_docx(filepath: str) -> str:
    """Parse Word 2007+ (.docx) file."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n\n'.join(paragraphs)
    except ImportError:
        raise ImportError(
            "python-docx required for .docx files. "
            "Install with: pip install python-docx"
        )


def parse_doc(filepath: str) -> str:
    """Parse legacy Word (.doc) file. Best-effort using antiword or textract."""
    # Try antiword first (common on Linux)
    try:
        import subprocess
        result = subprocess.run(
            ['antiword', filepath],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    # Try textract as fallback
    try:
        import textract
        return textract.process(filepath).decode('utf-8')
    except ImportError:
        pass
    
    raise ValueError(
        "Cannot parse .doc files. Options:\n"
        "1. Install antiword: apt install antiword\n"
        "2. Install textract: pip install textract\n"
        "3. Convert to .docx and re-upload"
    )


def parse_pdf(filepath: str) -> str:
    """Parse PDF file (text-based, not scanned images)."""
    try:
        import pypdf
        reader = pypdf.PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = '\n\n'.join(text_parts)
        
        if len(full_text.strip()) < 100:
            raise ValueError(
                "PDF appears to be scanned/image-based. "
                "Please provide a text-based PDF or paste the text directly."
            )
        
        return full_text
    except ImportError:
        raise ImportError(
            "pypdf required for .pdf files. "
            "Install with: pip install pypdf"
        )


def parse_rtf(filepath: str) -> str:
    """Parse RTF file."""
    try:
        from striprtf.striprtf import rtf_to_text
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            rtf_content = f.read()
        return rtf_to_text(rtf_content)
    except ImportError:
        raise ImportError(
            "striprtf required for .rtf files. "
            "Install with: pip install striprtf"
        )


# Format handlers
PARSERS = {
    '.txt': parse_txt,
    '.md': parse_txt,
    '.markdown': parse_txt,
    '.docx': parse_docx,
    '.doc': parse_doc,
    '.pdf': parse_pdf,
    '.rtf': parse_rtf,
}


def parse_file(filepath: str) -> ParsedDocument:
    """
    Parse a document file and extract text.
    
    Args:
        filepath: Path to the document
        
    Returns:
        ParsedDocument with extracted text and metadata
    """
    path = Path(filepath)
    
    if not path.exists():
        return ParsedDocument(
            filename=path.name,
            format='unknown',
            text='',
            word_count=0,
            success=False,
            error=f"File not found: {filepath}"
        )
    
    ext = path.suffix.lower()
    
    if ext not in PARSERS:
        return ParsedDocument(
            filename=path.name,
            format=ext,
            text='',
            word_count=0,
            success=False,
            error=f"Unsupported format: {ext}. Supported: {', '.join(PARSERS.keys())}"
        )
    
    try:
        text = PARSERS[ext](filepath)
        # Clean up text
        text = re.sub(r'\n{3,}', '\n\n', text)  # Normalize line breaks
        text = text.strip()
        word_count = len(text.split())
        
        return ParsedDocument(
            filename=path.name,
            format=ext,
            text=text,
            word_count=word_count,
            success=True
        )
    except Exception as e:
        return ParsedDocument(
            filename=path.name,
            format=ext,
            text='',
            word_count=0,
            success=False,
            error=str(e)
        )


def parse_files(filepaths: List[str]) -> Tuple[List[str], List[str]]:
    """
    Parse multiple files and return texts + errors.
    
    Args:
        filepaths: List of file paths
        
    Returns:
        Tuple of (successful_texts, error_messages)
    """
    texts = []
    errors = []
    
    for filepath in filepaths:
        result = parse_file(filepath)
        if result.success:
            if result.word_count < 100:
                errors.append(
                    f"⚠️ {result.filename}: Only {result.word_count} words. "
                    f"Need 500+ for good voice learning."
                )
            texts.append(result.text)
        else:
            errors.append(f"❌ {result.filename}: {result.error}")
    
    return texts, errors


def check_dependencies() -> dict:
    """Check which format dependencies are installed."""
    deps = {
        '.docx': False,
        '.pdf': False,
        '.rtf': False,
        '.doc': False,
    }
    
    try:
        import docx
        deps['.docx'] = True
    except ImportError:
        pass
    
    try:
        import pypdf
        deps['.pdf'] = True
    except ImportError:
        pass
    
    try:
        from striprtf.striprtf import rtf_to_text
        deps['.rtf'] = True
    except ImportError:
        pass
    
    try:
        import subprocess
        result = subprocess.run(['which', 'antiword'], capture_output=True)
        deps['.doc'] = result.returncode == 0
    except:
        pass
    
    return deps


# CLI interface
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python file_parser.py <file1> [file2] ...")
        print("\nChecking dependencies...")
        deps = check_dependencies()
        for fmt, installed in deps.items():
            status = "✅" if installed else "❌"
            print(f"  {fmt}: {status}")
        sys.exit(0)
    
    texts, errors = parse_files(sys.argv[1:])
    
    if errors:
        print("Errors:")
        for err in errors:
            print(f"  {err}")
        print()
    
    print(f"Successfully parsed {len(texts)} file(s)")
    for i, text in enumerate(texts):
        word_count = len(text.split())
        print(f"  File {i+1}: {word_count} words")
